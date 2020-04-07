#!python3

import pyinputplus as pyip
from roll import Roll
from player import Player
from pathlib import Path
from datetime import datetime
from docx2pdf import convert
import docx, os, time, sys


class Yahtzee:
    def __init__(self):

        # reference dicts and lists for calculating scores
        self._scoreDictReferenceValues = {
                'ones': 1,'twos': 2, 'threes': 3,
                'fours': 4, 'fives': 5, 'sixes': 6,
                'full house': 25, 'small straight': 30, 'large straight': 35,
                'yahtzee': 50, 'yahtzee bonus': 50,
                }

        # singles reference options when validating CHECK TOP SCORE
        self._singlesOptions = [
                'ones', 'twos', 'threes',
                'fours', 'fives', 'sixes'
                ]
        # player name strings
        self._playersNames = []

        # lists of class instances
        self._playersList = []
        self._rollsList = []

        # other objects
        self._numberOfPlayers = 0
        self._gameOver = False
        self._gameCounter = 0
        self._rankingDict = {}

    def getNumberOfPlayers(self):
        '''
        Gets the number of players (1 to 4).
        '''
        self._numberOfPlayers = pyip.inputInt(prompt='\nEnter number of players (1 to 4):\n', min=1, max=4)

    def getPlayerNames(self):
        '''
        Gets player names for number of players.
        '''
        for i in range(self._numberOfPlayers):
            self._playersNames.append(pyip.inputStr(prompt=f'\nEnter name of player {i+1}:\n'))

    def createPlayersList(self):
        '''
        Creates playersList of player instances (1 to 4)
        '''
        for playerName in self._playersNames:
            self._playersList.append(Player(playerName))

    def createRollsList(self):
        '''
        Creates playersList of player instances (1 to 4)
        '''
        for playerName in self._playersNames:
            self._rollsList.append(Roll(playerName))

    def sortRankingDict(self):
        '''
        Gets ranking dict of player and grand total score
        '''
        for j, player in enumerate(self._playersList):
            rankingName, rankingScore = self._playersList[j].getNameAndGrandTotalScore()
            self._rankingDict[rankingName] = rankingScore

        # reverse sort ranking dict by grand total
        self._rankingDict = sorted(self._rankingDict.items(), key=lambda x: x[1], reverse=True)

    def resetPlayerScores(self):
        '''
        Resets scores in all Player class instances for next game
        '''
        for j, player in enumerate(self._playersList):
            self._playersList[j].resetAllScores()

    def play(self):
        '''
        Main gameplay for yahtzee
        '''

        # get players count, player names, create instances of Player and Roll
        print('\nWELCOME TO YAHTZEE!')
        self.getNumberOfPlayers()
        self.getPlayerNames()
        self.createPlayersList()
        self.createRollsList()

        # GAME LOOP
        while self._gameOver is False:

            print(f"\nLET'S PLAY! GAME {self._gameCounter+1}")

            # ROUND LOOP (arbitrarily refs len of first instance of Player)
            for i, k in enumerate(self._playersList[0]._scoreDict):

                # PLAYER TURN
                for j, player in enumerate(self._playersList):

                    # skip final round if only yahtzee bonus and yahtzee != 50
                    if i == 13 and self._playersList[j]._scoreDict['yahtzee bonus'] == False and self._playersList[j]._scoreDict['yahtzee'] != 50:
                        self._playersList[j]._scoreDict['yahtzee bonus'] = 0
                        print("\nAutomatically score 0 for 'yahtzee bonus'...")

                    else:
                        # print current scores and totals before rolling
                        print(f'\n{self._playersList[j].name.upper()} YOUR TURN. ROUND: {i+1}')

                        print('-'*21)
                        print('ROLL SCORES'.rjust(16))
                        self._playersList[j].printStackedScoreDict()

                        print('-'*21)
                        print('TOP SCORE BONUS'.rjust(19))
                        print(self._playersList[j].getTopScore())
                        print(self._playersList[j].getTopBonusScore())

                        print('-'*21)
                        print('TOTAL SCORES'.rjust(19))
                        print(self._playersList[j].getTotalTopScore())
                        print(self._playersList[j].getTotalBottomScore())

                        print('-'*21)
                        print(f"{self._playersList[j].getGrandTotalScore()}\n")

                        # first roll
                        self._rollsList[j].rollDice()
                        print(f'{self._playersList[j].name.upper()}', end='')
                        keepFirstRoll = self._rollsList[j].keepDice()

                        # second roll
                        self._rollsList[j].reRollDice(keepFirstRoll)
                        print(f'{self._playersList[j].name.upper()}', end='')
                        keepSecondRoll = self._rollsList[j].keepDice()

                        # third roll
                        finalRoll = self._rollsList[j].finalRollDice(keepSecondRoll)

                        # select score to check final roll against
                        scoreSelected = self._playersList[j].selectScore(finalRoll)

                        # CHECK TOP or BOTTOM SCORE per score option selected
                        # CHECK TOP SCORE and increment score
                        if scoreSelected in self._singlesOptions:
                            score = self._rollsList[j].checkSingles(finalRoll, self._scoreDictReferenceValues[scoreSelected])

                            # incremenet score option, top score, top total, grand total
                            self._playersList[j]._scoreDict[scoreSelected] += score
                            self._playersList[j]._topScore += score
                            self._playersList[j]._totalTopScore += score
                            self._playersList[j]._grandTotalScore += score

                            # check top bonus, increment top total and grand total
                            self._playersList[j].addTopBonusScore()
                            self._playersList[j]._totalTopScore += self._playersList[j]._topBonusScore
                            self._playersList[j]._grandTotalScore += self._playersList[j]._topBonusScore

                        # CHECK BOTTOM SCORE and increment score
                        else:
                            if scoreSelected == 'three of a kind':
                                score = self._rollsList[j].checkThreeOfAKind(finalRoll)

                            elif scoreSelected == 'four of a kind':
                                score = self._rollsList[j].checkFourOfAKind(finalRoll)

                            elif scoreSelected == 'full house':
                                score = self._rollsList[j].checkFullHouse(finalRoll)

                            elif scoreSelected == 'small straight':
                                score = self._rollsList[j].checkSmallStraight(finalRoll)

                            elif scoreSelected == 'large straight':
                                score = self._rollsList[j].checkLargeStraight(finalRoll)

                            elif scoreSelected == 'yahtzee':
                                score = self._rollsList[j].checkYahtzee(finalRoll)

                            elif scoreSelected == 'chance':
                                score = self._rollsList[j].addChance(finalRoll)

                            elif scoreSelected == 'yahtzee bonus':
                                score = self._rollsList[j].checkYahtzeeBonus(finalRoll)

                                # player cannot score yahztee bonus 50 if no yahtzee
                                if self._playersList[j]._scoreDict['yahtzee'] != 50:
                                    score = 0

                            # increment round, total bottom, and grand total scores
                            self._playersList[j]._scoreDict[scoreSelected] += score
                            self._playersList[j]._totalBottomScore += score
                            self._playersList[j]._grandTotalScore += score

                        # print grand total score for end of player turn
                        print(f"\n{self._playersList[j].name.upper()} GRAND TOTAL: {self._playersList[j]._grandTotalScore}")
                        print("-"*21)

            # END OF ROUND PRINT ACTIONS

            # create ranking dict for the round
            self.sortRankingDict()

            # print rankings for the round
            print('\nFINAL SCORES')
            print('-'*12)
            for k, v in enumerate(self._rankingDict):
                print(f"{k+1} {v[0]}: {v[1]}")
            print('\n')

            # END OF ROUND FILE I/O

            # create YahtzeeScores directory
            os.makedirs(Path.cwd() / 'YahtzeeScores', exist_ok=True)

            # CREATE TEXT FILE

            # create TextFiles Directory
            os.makedirs(Path.cwd() / 'YahtzeeScores/TextFiles', exist_ok=True)
            textFileDirStr = str(Path.cwd() / 'YahtzeeScores/TextFiles')

            # create text file filename with datetime and game number
            dateToday = datetime.today().strftime('%Y-%m-%d-%H:%M:%S')
            textFilename = f"{dateToday}Game{self._gameCounter+1}.txt"

            # write scores to text file
            with open(f'{textFileDirStr}/{textFilename}', 'w') as f:
                f.write(f'YAHTZEE GAME {self._gameCounter+1}\n')
                f.write('FINAL RANKINGS\n')

                # write ranking of all players to file
                f.write(f"{'-'*21}")
                for k, v in enumerate(self._rankingDict):
                    f.write(f"\n{v[0]}: {v[1]}")
                f.write(f"\n{'-'*21}\n")

                # write each player score dict and total scores to file
                for j, player in enumerate(self._playersList):
                    f.write(f"\n{'-'*21}")
                    f.write(f"\n{'-'*21}")
                    f.write(f"\n{' '*2}{self._playersList[j].name.upper()} FINAL SCORES\n")

                    f.write(f"\n{'ROLL SCORES'.rjust(16)}")
                    outputScoreDict = self._playersList[j].getScoreDict()
                    for i, k in enumerate(outputScoreDict):
                        f.write(f"\n{k.rjust(15)}: {outputScoreDict[k]}")

                    f.write(f"\n{'-'*21}\n")
                    f.write(f"{'TOP SCORE BONUS'.rjust(19)}\n")
                    f.write(f"{self._playersList[j].getTopScore()}\n".rjust(20))
                    f.write(f"{self._playersList[j].getTopBonusScore()}\n".rjust(20))

                    f.write(f"\n{'TOTAL SCORES'.rjust(19)}\n")
                    f.write(f"{self._playersList[j].getTotalTopScore()}\n".rjust(20))
                    f.write(f"{self._playersList[j].getTotalBottomScore()}\n".rjust(20))

                    f.write(f"\n{'-'*21}\n")
                    f.write(f"{self._playersList[j].getGrandTotalScore()}\n".rjust(20))
                    f.write('\n')
            print("\nSaved scores textfile in: 'YahtzeeScores/TextFiles/'...")

            # CREATE WORD FILE

            # create Docx Directory
            os.makedirs(Path.cwd() / 'YahtzeeScores/DocxFiles/', exist_ok=True)
            docxFileDirStr = str(Path.cwd() / 'YahtzeeScores/DocxFiles/')

            # create docx file filename with datetime and game number
            docxFilename = f"{dateToday}Game{self._gameCounter+1}.docx"

            # open blank Document object
            doc = docx.Document()
            doc.add_paragraph(f'YAHTZEE GAME {self._gameCounter+1}', 'Title')
            doc.paragraphs[0].runs[0].add_break()

            # add picture of yahtzee game to document
            doc.add_picture(str(Path.cwd() / 'yahtzeePicture.jpg'))

            doc.add_heading('FINAL RANKINGS', 1)
            for k, v in enumerate(self._rankingDict):
                doc.add_paragraph(f"{v[0]}: {v[1]}")

            # add page break after rankings
            paraObjRankings = doc.add_paragraph('   ')
            paraObjRankings.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

            # write each player score dict and total scores to file
            doc.add_heading('PLAYER SCORES AND TOTALS', 1)
            for j, player in enumerate(self._playersList):

                # write player name as header
                doc.add_heading(f"{self._playersList[j].name.upper()}", 2)

                # write player roll scores for each scoring option
                doc.add_heading('ROLL SCORES', 3)
                outputScoreDict = self._playersList[j].getScoreDict()
                for i, k in enumerate(outputScoreDict):
                    doc.add_paragraph(f"{k}: {outputScoreDict[k]}")

                # write top score and bonus
                doc.add_heading('TOP SCORE BONUS', 3)
                doc.add_paragraph(f"{self._playersList[j].getTopScore()}")
                doc.add_paragraph(f"{self._playersList[j].getTopBonusScore()}")

                # write total scores and grand total
                doc.add_heading('TOTAL SCORES', 3)
                doc.add_paragraph(f"{self._playersList[j].getTotalTopScore()}")
                doc.add_paragraph(f"{self._playersList[j].getTotalBottomScore()}")
                paraObjGT = doc.add_paragraph(f"{self._playersList[j].getGrandTotalScore()}")

                # add pagebreak before writing next player scores to docx
                if j != (len(self._playersList)-1):
                    paraObjGT.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

            # save Document object as docxFilename
            doc.save(f"{docxFileDirStr}/{docxFilename}")
            print("\nSaved .docx scores file in: 'YahtzeeScores/DocxFiles/'...")

            # CONVERT TO PDF

            # create PDF Directory
            os.makedirs(Path.cwd() / 'YahtzeeScores/pdfFiles/', exist_ok=True)
            pdfFileDirStr = str(Path.cwd() / 'YahtzeeScores/pdfFiles/')

            # convert docx to pdf
            convert(f"{docxFileDirStr}/{docxFilename}", f"{pdfFileDirStr}/{docxFilename[:-5]}.pdf")
            print("\nSaved pdf scores file in: 'YahtzeeScores/pdfFiles/'...")

            # END OF GAME ACTIONS

            # reset each Player class instance scoring dict and total scores
            print('\nResetting dice for next round...')
            time.sleep(1)
            self.resetPlayerScores()

            # increment game counter, if three games end match
            self._gameCounter += 1

            if self._gameCounter == 3:
                print('GAME OVER')
                self._gameOver = True

        # exit game
        sys.exit()
