#!python3

from pathlib import Path
from docx2pdf import convert
import docx, os


def createFileioDirectory():
    '''
    Creates YahtzeeScores/ folder in cwd.
    '''
    os.makedirs(Path.cwd() / 'YahtzeeScores', exist_ok=True)


def printFileioConfirmation(fileDirStr, fileName):
    '''
    Prints confirmation message when file creation completed.

    :param fileDirStr: textfile directory as string
    :param fileName: textfile basename
    '''
    print(f"\nSaved file: '{fileDirStr}/{fileName}'")


class TextFile:
    '''
    Objects instantiated by the :class:`TextFile <Textfile>` can be called to create a text file of players and scores.
    '''
    def __init__(self):
        # instance objects
        self._textFileDirStr = ''
        self._textFilename = ''

    def __repr__(self):
        return (f"{self.__class__.__name__}()")

    def createTextFileDir(self):
        '''
        Create TextFiles folder.
        '''
        os.makedirs(Path.cwd() / 'YahtzeeScores/TextFiles', exist_ok=True)
        self._textFileDirStr = str(Path.cwd() / 'YahtzeeScores/TextFiles')

    def createTextFilename(self, gameCounter, dateTimeToday):
        '''
        Create text file filename with datetime and game number.

        :param gameCounter: integer count of games played.
        :param dateTimeToday: date str to standardize output file basename.
        '''
        self._textFilename = f"{dateTimeToday}Game{gameCounter+1}.txt"

    def writeTextFile(self, gameCounter, playersList, rankingDict):
        '''
        Writes players scores to text file.

        :param gameCounter: integer count of games played.
        :param playersList: list of Player class instances.
        :param rankingDict: ranking of players and grand total scores.
        '''
        with open(f'{self._textFileDirStr}/{self._textFilename}', 'w') as f:
            f.write(f'YAHTZEE GAME {gameCounter+1}\n')
            f.write('FINAL RANKINGS\n')

            # write ranking of all players to file
            f.write(f"{'-'*21}")
            for k, v in enumerate(rankingDict):
                f.write(f"\n{v[0]}: {v[1]}")
            f.write(f"\n{'-'*21}\n")

            # write each player score dict and total scores to file
            for j, player in enumerate(playersList):
                f.write(f"\n{'-'*21}")
                f.write(f"\n{'-'*21}")
                f.write(f"\n{' '*2}{playersList[j].name.upper()} FINAL SCORES\n")

                f.write(f"\n{'ROLL SCORES'.rjust(16)}")
                outputScoreDict = playersList[j].getScoreDict()
                for i, k in enumerate(outputScoreDict):
                    f.write(f"\n{k.rjust(15)}: {outputScoreDict[k]}")

                f.write(f"\n{'-'*21}\n")
                f.write(f"{'TOP SCORE BONUS'.rjust(19)}\n")
                f.write(f"{playersList[j].getTopScore()}\n".rjust(20))
                f.write(f"{playersList[j].getTopBonusScore()}\n".rjust(20))

                f.write(f"\n{'TOTAL SCORES'.rjust(19)}\n")
                f.write(f"{playersList[j].getTotalTopScore()}\n".rjust(20))
                f.write(f"{playersList[j].getTotalBottomScore()}\n".rjust(20))

                f.write(f"{'-'*21}\n")
                f.write(f"{playersList[j].getGrandTotalScore()}".rjust(20))
                f.write('\n')

            # print file creation confirmation
            printFileioConfirmation(self._textFileDirStr, self._textFilename)


class DocxFile:
    '''
    Objects instantiated by the :class:`DocxFile <DocxFile>` can be called to create a docx file of players and scores.
    '''
    def __init__(self):
        # instance objects
        self._docxFileDirStr = ''
        self._docxFilename = ''

    def __repr__(self):
        return (f"{self.__class__.__name__}()")

    def createDocxFileDir(self):
        '''
        Create DocxFiles folder.
        '''
        os.makedirs(Path.cwd() / 'YahtzeeScores/DocxFiles', exist_ok=True)
        self._docxFileDirStr = str(Path.cwd() / 'YahtzeeScores/DocxFiles')

    def createDocxFilename(self, gameCounter, dateTimeToday):
        '''
        Create docx file filename with datetime and game number.

        :param gameCounter: integer count of games played.
        :param dateTimeToday: date str to standardize output file basename.
        '''
        self._docxFilename = f"{dateTimeToday}Game{gameCounter+1}.docx"

    def writeDocxFile(self, gameCounter, playersList, rankingDict):
        '''
        Writes players scores to docx file.

        :param gameCounter: integer count of games played.
        :param playersList: list of Player class instances.
        :param rankingDict: ranking of players and grand total scores.
        '''
        # open blank Document object
        doc = docx.Document()
        doc.add_paragraph(f'YAHTZEE GAME {gameCounter+1}', 'Title')
        doc.paragraphs[0].runs[0].add_break()

        # add picture of yahtzee game to document
        doc.add_picture(str(Path.cwd() / 'yahtzeePicture.jpg'))

        doc.add_heading('FINAL RANKINGS', 1)
        for k, v in enumerate(rankingDict):
            doc.add_paragraph(f"{v[0]}: {v[1]}")

        # add page break after rankings
        paraObjRankings = doc.add_paragraph('   ')
        paraObjRankings.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

        # write each player score dict and total scores to file
        doc.add_heading('PLAYER SCORES AND TOTALS', 1)
        for j, player in enumerate(playersList):

            # write player name as header
            doc.add_heading(f"{playersList[j].name.upper()}", 2)

            # write player roll scores for each scoring option
            doc.add_heading('ROLL SCORES', 3)
            outputScoreDict = playersList[j].getScoreDict()
            for i, k in enumerate(outputScoreDict):
                doc.add_paragraph(f"{k}: {outputScoreDict[k]}")

            # write top score and bonus
            doc.add_heading('TOP SCORE BONUS', 3)
            doc.add_paragraph(f"{playersList[j].getTopScore()}")
            doc.add_paragraph(f"{playersList[j].getTopBonusScore()}")

            # write total scores and grand total
            doc.add_heading('TOTAL SCORES', 3)
            doc.add_paragraph(f"{playersList[j].getTotalTopScore()}")
            doc.add_paragraph(f"{playersList[j].getTotalBottomScore()}")
            paraObjGT = doc.add_paragraph(f"{playersList[j].getGrandTotalScore()}")

            # add pagebreak before writing next player scores to docx
            if j != (len(playersList)-1):
                paraObjGT.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

        # save Document object as docxFilename
        doc.save(f"{self._docxFileDirStr}/{self._docxFilename}")

        # print file creation confirmation
        printFileioConfirmation(self._docxFileDirStr, self._docxFilename)


class PdfFile:
    '''
    Objects instantiated by the :class:`DocxFile <DocxFile>` can be called to convert a docx file to a pdf file
    '''
    def __init__(self):
        self._pdfFileDirStr = ''
        self._pdfFilename = ''

    def createPdfFileDir(self):
        '''
        Create PDF files folder.
        '''
        os.makedirs(Path.cwd() / 'YahtzeeScores/pdfFiles/', exist_ok=True)
        self._pdfFileDirStr = str(Path.cwd() / 'YahtzeeScores/pdfFiles/')

    def createPdfFilename(self, gameCounter, dateTimeToday):
        '''
        Create pdf file filename with datetime and game number.

        :param gameCounter: integer count of games played.
        :param dateTimeToday: date str to standardize output file basename.
        '''
        self._pdfFilename = f"{dateTimeToday}Game{gameCounter+1}.pdf"

    def convertDocxToPdf(self, docxFileDirStr, docxFilename):
        '''
        Converts Docx file to Pdf file
        '''
        convert(f"{docxFileDirStr}/{docxFilename}", f"{self._pdfFileDirStr}/{self._pdfFilename}")

        # print file convert confirmation
        printFileioConfirmation(self._pdfFileDirStr, self._pdfFilename)
