#!python3

"""
This module controls file input/output of yahtzee score files.

github.com/pmacking/fileio.py
"""

import os
from pathlib import Path
import docx


def print_fileio_confirmation(file_dir_str, file_name):
    """
    Prints confirmation message when file creation completed.

    :param file_dir_str: textfile directory as string
    :param file_name: textfile basename
    """
    print(f"\nSaved file: '{file_dir_str}/{file_name}'")


class FileWriter:
    """
    Objects instantiated by :class: `FileWriter <FileWriter>` can be called as
    a factory to write file as output per file formats.
    """
    # attributes specified for docx only to enable passing to pdf conversion
    def __init__(self):
        self.docxfile_dir_str = ''
        self.docx_filename = ''

    def __repr__(self):
        return f"{self.__class__.__name__}()"

    def write_file(self, datetime_today, game_counter,
                   players_list, ranking_dict, file_formats):
        """
        The write_file method creates instances of each file_format, including
        creating Path, filename, and writing or converting files.

        :param datetime_today: The datetime today as string.
        :param game_counter: The count of the game as int.
        :param players_list: The list of player instances.
        :param ranking_dict: The ranking dictionary for the current round.
        :param file_formats: The file formats to write as list.
        """

        # create class and write file for each format in file_formats
        for file_format in file_formats:
            if file_format == 'txt':
                textfile = TextFile()

                # create textfile directory
                textfile.create_textfile_dir()

                # create textfile basename
                textfile.create_textfile_name(game_counter, datetime_today)

                # write textfile
                textfile.write_text_file(game_counter, players_list,
                                         ranking_dict)

            elif file_format == 'docx':
                docxfile = DocxFile()

                # create docx directory and set locally for pdf convert
                self.docxfile_dir_str = docxfile.create_docxfile_dir()

                # create docx basename and set locally for pdf convert
                self.docx_filename = docxfile.create_docx_filename(
                                        game_counter, datetime_today)

                # write docxfile
                docxfile.write_docxfile(game_counter, players_list,
                                        ranking_dict)

            # REMOVED: ISSUES WITH MSFT WORD and NOT SUPPORTED ON LINUX
            # elif file_format == 'pdf':
            #     # PDF instance in fileio.py
            #     pdffile = PdfFile()

            #     # create pdf file directory
            #     pdffile.create_pdffile_dir()

            #     # create pdf file basename
            #     pdffile.create_pdf_filename(game_counter, datetime_today)

            #     # convert docx to pdf
            #     pdffile.convertDocxToPdf(self.docxfile_dir_str,
            #                              self.docx_filename)


class TextFile:
    """
    Objects instantiated by the :class:`TextFile <Textfile>` can be called to
    create a textfile of players and scores.
    """
    def __init__(self):
        self.relative_path = 'data/Textfiles'
        self.textfile_dir_str = ''
        self.textfile_name = ''

    def __repr__(self):
        return (f"{self.__class__.__name__}("
                f"{self.textfile_dir_str}, {self.textfile_name})")

    def create_textfile_dir(self):
        """
        Create Textfiles folder.
        """
        os.makedirs(Path.cwd() / f'{self.relative_path}', exist_ok=True)
        self.textfile_dir_str = str(Path.cwd() / f'{self.relative_path}')

    def create_textfile_name(self, game_counter, datetime_today):
        """
        Create textfile filename with datetime and game number.

        :param game_counter: integer count of games played.
        :param datetime_today: date str to standardize output file basename.
        """
        self.textfile_name = f"{datetime_today}Game{game_counter + 1}.txt"

    def write_text_file(self, game_counter, players_list, ranking_dict):
        """
        Writes players scores to textfile.

        :param game_counter: integer count of games played.
        :param players_list: list of Player class instances.
        :param ranking_dict: ranking of players and grand total scores.
        """
        with open(f'{self.textfile_dir_str}/{self.textfile_name}', 'w') as f:
            f.write(f'YAHTZEE GAME {game_counter+1}\n')
            f.write('FINAL RANKINGS\n')

            # write ranking of all players to file
            f.write(f"{'-'*21}")
            for k, v in enumerate(ranking_dict):
                f.write(f"\n{v[0]}: {v[1]}")
            f.write(f"\n{'-'*21}\n")

            # enumerate players and write scores to file
            for j, player in enumerate(players_list):
                f.write(f"\n{'-'*21}")
                f.write(f"\n{'-'*21}")
                f.write(f"\n{' '*2}{players_list[j].name.upper()} "
                        f"FINAL SCORES\n")

                f.write(f"\n{'ROLL SCORES'.rjust(16)}")

                # write player's score dictionary to file
                output_score_dict = players_list[j].get_score_dict()
                for i, k in enumerate(output_score_dict):
                    f.write(f"\n{k.rjust(15)}: {output_score_dict[k]}")

                # write top, total, and grand total scores to file
                f.write(f"\n{'-'*21}\n")
                f.write(f"{'TOP SCORE BONUS'.rjust(19)}\n")
                f.write(f"{players_list[j].get_top_score()}\n".rjust(20))
                f.write(f"{players_list[j].get_top_bonus_score()}\n".rjust(20))

                f.write(f"\n{'TOTAL SCORES'.rjust(19)}\n")
                f.write(f"{players_list[j].get_total_top_score()}\n".rjust(20))
                f.write(f"{players_list[j].get_total_bottom_score()}\n".rjust(20))

                f.write(f"{'-'*21}\n")
                f.write(f"{players_list[j].get_grand_total_score()}".rjust(20))
                f.write('\n')

            # print file creation confirmation
            print_fileio_confirmation(self.textfile_dir_str,
                                      self.textfile_name)


class DocxFile:
    """
    Objects instantiated by the :class:`DocxFile <DocxFile>` can be called to
    create a docx file of players and scores.
    """
    def __init__(self):
        self.relative_path = 'data/Docxfiles'
        self.docxfile_dir_str = ''
        self.docx_filename = ''

    def __repr__(self):
        return (f"{self.__class__.__name__}("
                f"{self.docxfile_dir_str}, {self.docx_filename})")

    def create_docxfile_dir(self):
        """
        Create Docxfiles folder.

        :rtype: string
        :return: The docx directory Path.
        """
        os.makedirs(Path.cwd() / f'{self.relative_path}', exist_ok=True)
        self.docxfile_dir_str = str(Path.cwd() / f'{self.relative_path}')

        return self.docxfile_dir_str

    def create_docx_filename(self, game_counter, datetime_today):
        """
        Create docx filename with datetime and game number.

        :param game_counter: integer count of games played.
        :param datetime_today: date str to standardize output file basename.

        :return: The docx filename as string.
        """
        self.docx_filename = f"{datetime_today}Game{game_counter+1}.docx"

        return self.docx_filename

    def write_docxfile(self, game_counter, players_list, ranking_dict):
        """
        Writes players scores to docx file.

        :param game_counter: integer count of games played.
        :param players_list: list of Player class instances.
        :param ranking_dict: ranking of players and grand total scores.
        """
        # open blank Document object
        doc = docx.Document()
        doc.add_paragraph(f'YAHTZEE GAME {game_counter + 1}', 'Title')
        doc.paragraphs[0].runs[0].add_break()

        # add picture of yahtzee game to document
        doc.add_picture(
            str(Path.cwd() / 'yahtzee/resources/yahtzeePicture.jpg'))

        doc.add_heading('FINAL RANKINGS', 1)
        for k, v in enumerate(ranking_dict):
            doc.add_paragraph(f"{v[0]}: {v[1]}")

        # add page break after rankings
        para_obj_rankings = doc.add_paragraph('   ')
        para_obj_rankings.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

        # write each player score dict and total scores to file
        doc.add_heading('PLAYER SCORES AND TOTALS', 1)
        for j, player in enumerate(players_list):

            # write player name as header
            doc.add_heading(f"{players_list[j].name.upper()}", 2)

            # write scores for each scoring option
            doc.add_heading('ROLL SCORES', 3)
            output_score_dict = players_list[j].get_score_dict()
            for i, k in enumerate(output_score_dict):
                doc.add_paragraph(f"{k}: {output_score_dict[k]}")

            # write top score and bonus
            doc.add_heading('TOP SCORE BONUS', 3)
            doc.add_paragraph(f"{players_list[j].get_top_score()}")
            doc.add_paragraph(f"{players_list[j].get_top_bonus_score()}")

            # write total scores and grand total
            doc.add_heading('TOTAL SCORES', 3)
            doc.add_paragraph(f"{players_list[j].get_total_top_score()}")
            doc.add_paragraph(f"{players_list[j].get_total_bottom_score()}")
            para_obj_grand_total = doc.add_paragraph(
                            f"{players_list[j].get_grand_total_score()}")

            # add pagebreak before writing next player scores to docx
            if j != (len(players_list)-1):
                para_obj_grand_total.runs[0].add_break(
                    docx.enum.text.WD_BREAK.PAGE)

        # save Document object as docx filename
        doc.save(f"{self.docxfile_dir_str}/{self.docx_filename}")

        # print file creation confirmation
        print_fileio_confirmation(self.docxfile_dir_str, self.docx_filename)


# class PdfFile:
#     """
#     Objects instantiated by the :class:`DocxFile <DocxFile>` can be called to
#     convert a docx file to a pdf file
#     """
#     def __init__(self):
#         self.pdffile_dir_str = ''
#         self.pdf_filename = ''

#     def __repr__(self):
#         return (f"{self.__class__.__name__}("
#                 f"{self.pdffile_dir_str}, {self.pdf_filename})")

#     def create_pdffile_dir(self):
#         """
#         Create PDF files folder.
#         """
#         os.makedirs(Path.cwd() / 'data/pdf_files/', exist_ok=True)
#         self.pdffile_dir_str = str(Path.cwd() / 'data/pdf_files/')

#     def create_pdf_filename(self, game_counter, datetime_today):
#         """
#         Create pdf filename with datetime and game number.

#         :param game_counter: integer count of games played.
#         :param datetime_today: date str to standardize output file basename.
#         """
#         self.pdf_filename = f"{datetime_today}Game{game_counter + 1}.pdf"

#     def convertDocxToPdf(self, docxfile_dir_str, docx filename):
#         """
#         Converts Docx file to Pdf file
#         """
#         convert(f"{docxfile_dir_str}/{docx filename}",
#                 f"{self.pdffile_dir_str}/{self.pdf_filename}")

#         # print file convert confirmation
#         print_fileio_confirmation(self.pdffile_dir_str, self.pdf_filename)
